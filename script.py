import tkinter as tk
from tkinter import messagebox
from tkinter import ttk
import os
import qrcode
from openpyxl import load_workbook # type: ignore
from docx import Document # type: ignore
import PyPDF2
import time
def convertFile():
    fileName=file.get().strip()
    nameQr=name.get().strip()
    user_path= os.path.expanduser('~')
    path = os.path.join(os.path.expanduser("~"), "Desktop")
    qr_code_folder = os.path.join(path, "qr_code")
    os.makedirs(qr_code_folder, exist_ok=True)
    ok=False
    progress['value']=0
    window.update_idletasks()
    try:
        
        for root,dirs,files in os.walk(user_path):
            if fileName in files :
                full_file_path = os.path.join(root, fileName) 
                index=fileName.find('.')
                extension=fileName[index+1:]
                progress['value']=20
                window.update_idletasks()
                time.sleep(1)
                if(extension=='txt'):
                    
                    with open(full_file_path, 'r', encoding='utf-8') as file_content:
                        data =file_content.read()
                    progress['value']=50
                    window.update_idletasks()
                    time.sleep(1)
                    qr=qrcode.make(data)
                    qr.save(os.path.join(qr_code_folder, f"{nameQr}.png"))
                    progress['value']=100
                    window.update_idletasks()
                    time.sleep(1)
                    messagebox.showinfo("Succès", "QR code généré et enregistré avec succès")
                    ok=True
                if(extension=='docx'):
                    doc=Document(full_file_path)
                    progress['value']=50
                    window.update_idletasks()
                    time.sleep(1)
                    Document(full_file_path) 
                    data1 = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    qr=qrcode.make(data1)
                    qr.save(os.path.join(qr_code_folder, f"{nameQr}.png"))
                    progress['value']=100
                    window.update_idletasks()
                    time.sleep(1)
                    messagebox.showinfo("Succès", "QR code généré et enregistré avec succès")
                    ok=True
                if(extension=='pdf'):
                    with open(full_file_path, 'rb') as file_pdf:
                        read_pdf = PyPDF2.PdfReader(file_pdf)
                        content = ""
                        for page in range(len(read_pdf.pages)): 
                            content += read_pdf.pages[page].extract_text()
                    progress['value']=50
                    window.update_idletasks()
                    time.sleep(1)
                    qr=qrcode.make(content)
                    qr.save(os.path.join(qr_code_folder, f"{nameQr}.png"))
                    progress['value']=100
                    window.update_idletasks()
                    time.sleep(1)
                    messagebox.showinfo("Succès", "QR code généré et enregistré avec succès")
                    ok=True
                if(extension=='xlsx' or extension=='xls'):
                        wb = load_workbook(full_file_path) 
                        sheet = wb.active
                        data2 = "" 
                        for row in sheet.iter_rows(values_only=True): 
                             for cell in reversed(row):
                                 if(str(cell).lower()!='none'):
                                    row_data = "\t"+str(cell)
                                    data2 += f"{row_data}\n"
                        progress['value']=50
                        window.update_idletasks()
                        time.sleep(1)
                        qr=qrcode.make(data2)
                        qr.save(os.path.join(qr_code_folder, f"{nameQr}.png"))
                        progress['value']=100
                        window.update_idletasks()
                        time.sleep(1)
                        messagebox.showinfo("Succès", "QR code généré et enregistré avec succès")
                        ok=True
                    
    except Exception :
                messagebox.showerror("erreur","une erreur s'est produite !")
    if not ok :
      messagebox.showerror("Erreur","le fichier n'existe pas")
def convertLink():
    progress['value']=0
    window.update_idletasks()
    data = link.get().strip()
    name_value = name.get().strip()
    path = os.path.join(os.path.expanduser("~"), "Desktop")
    progress['value']=50
    window.update_idletasks()
    time.sleep(1)
    qr_code_folder = os.path.join(path, "qr_code")
    os.makedirs(qr_code_folder, exist_ok=True)
    qr = qrcode.make(data)
    qr.save(os.path.join(qr_code_folder, f"{name_value}.png"))
    progress['value']=100
    window.update_idletasks()
    time.sleep(1)
    messagebox.showinfo("Succès", "QR code généré et enregistré avec succès")
def convertFileFromSearch():
            fileName=file_search.get().strip()
            name1=name_search.get().strip()
            name2=surname_search.get().strip()
            qrname=qr_n.get().strip()
            user_path= os.path.expanduser('~')
            path = os.path.join(os.path.expanduser("~"), "Desktop")
            qr_code_folder = os.path.join(path, "qr_code")
            os.makedirs(qr_code_folder, exist_ok=True)
            ok=False
            
            s=0
            progress['value']=0
            window.update_idletasks()
            try:
                
                for root,dirs,files in os.walk(user_path):
                    if fileName in files :
                        full_file_path = os.path.join(root, fileName) 
                        index=fileName.find('.')
                        extension=fileName[index+1:]
                        progress['value']=20
                        window.update_idletasks()
                        time.sleep(1)
                        if(extension=='xlsx' or extension=='xls'):
                                wb = load_workbook(full_file_path) 
                                sheet = wb.active
                                data2 = "" 
                                for row in sheet.iter_rows(values_only=True): 
                                    for cell in reversed(row):
                                        if(str(cell).strip().lower()==name1.lower()):
                                            
                                            
                                            s=1
                                        if(str(cell).strip().lower()==name2.lower() and s==1):
                                            s=2
                                            
                                    if(s==2):
                                        for cell in reversed(row):
                                          if(str(cell)!='none' and str(cell)!='None' ):

                                                data2 += "\t"+str(cell)
                                        progress['value']=50
                                        window.update_idletasks()
                                        time.sleep(1)
                                        print(data2)
                                        qr=qrcode.make(data2)
                                        qr.save(os.path.join(qr_code_folder, f"{qrname}.png"))
                                        progress['value']=100
                                        window.update_idletasks()
                                        time.sleep(1)
                                        messagebox.showinfo("Succès", "QR code généré et enregistré avec succès")
                                        ok=True
                                        break
                               
                                  
                                if(s!=2):
                                    messagebox.showerror("erreur","le nom ou prenom n'est pas trouvé !")
                                
                        else:
                                messagebox.showerror("erreur","il faut un fichier Excel")

            except Exception :
                        messagebox.showerror("erreur","une erreur s'est produite !")
            if not ok :
             messagebox.showerror("Erreur","le fichier n'existe pas")
#--------------------------------------------------------------------------------------
def openConvertFile():
    convertFile()
def openConvertLink(): 
    convertLink() 
def showAbout(): 
    messagebox.showinfo("À propos", "Application de Génération de QrCode: utilisez le champs fichier pour les fichiers locaux (.txt/.pdf/.docx/.xlsx/.xls),utilisez le champs lien pour une redirection à un fichier dans le oneDrive ou autre")
def openSearchWindow(): 
    search_window = tk.Toplevel(window) 
    search_window.title("Recherche de Fichier") 
    search_window.geometry("400x450") 
    search_window.configure(bg='black') 
    search_font_style = ("Helvetica", 12, "bold") 
    search_label = tk.Label(search_window, text="Donner le nom de fichier :", bg='black', fg='white', font=search_font_style) 
    search_label.pack(anchor="center", pady=10) 
    global file_search 
    file_search= tk.Entry(search_window, fg='white', bg='gray', font=search_font_style, insertbackground='white') 
    file_search.pack(anchor="center", pady=10, ipadx=10, ipady=5) 
    file_search.insert(0, "e.g: fichier.xlsx") 
    search_labeln = tk.Label(search_window, text="Donner le nom que vous voulez trouver :", bg='black', fg='white', font=search_font_style) 
    search_labeln.pack(anchor="center", pady=10) 
    global name_search
    name_search = tk.Entry(search_window, fg='white', bg='gray', font=search_font_style, insertbackground='white') 
    name_search.pack(anchor="center", pady=10, ipadx=10, ipady=5) 
    name_search.insert(0, "e.g: TOUZRI") 
    search_labelp = tk.Label(search_window, text="Donner le prénom que vous voulez trouver :", bg='black', fg='white', font=search_font_style) 
    search_labelp.pack(anchor="center", pady=10) 
    global surname_search
    surname_search = tk.Entry(search_window, fg='white', bg='gray', font=search_font_style, insertbackground='white') 
    surname_search.pack(anchor="center", pady=10, ipadx=10, ipady=5) 
    surname_search.insert(0, "e.g: YOUSSEF") 
    search_labelq= tk.Label(search_window, text="Donner le nom de Qrcode :", bg='black', fg='white', font=search_font_style) 
    search_labelq.pack(anchor="center", pady=10) 
    global qr_n
    qr_n = tk.Entry(search_window, fg='white', bg='gray', font=search_font_style, insertbackground='white') 
    qr_n.pack(anchor="center", pady=10, ipadx=10, ipady=5) 
    qr_n.insert(0, "e.g: les_donnees_en_qrcode") 
    search_button = tk.Button(search_window, text="Convertir les données", command=convertFileFromSearch, bg='purple', fg='white', font=search_font_style) 
    search_button.pack(pady=10, ipadx=10, ipady=5)
window = tk.Tk() 
window.title("Générateur de Qrcode") 
window.geometry("500x400") 
window.configure(bg='black') 
font_style = ("Helvetica", 14, "bold")
labelQrName = tk.Label(window, text="Le nom du Qrcode :", bg='black', fg='white', font=font_style) 
labelQrName.pack(anchor="center", pady=10) 
name = tk.Entry(window, fg='white', bg='gray', font=font_style, insertbackground='white') 
name.pack(anchor="center", pady=10, ipadx=10, ipady=5) 
name.insert(0, "e.g:myfirstqrcode") 
labelFile = tk.Label(window, text="Fichier :", bg='black', fg='white', font=font_style) 
labelFile.pack(anchor="center", pady=10) 
file = tk.Entry(window, fg='white', bg='gray', font=font_style, insertbackground='white') 
file.pack(anchor="center", pady=10, ipadx=10, ipady=5) 
file.insert(0, "e.g:fichier.txt") 
labelLink = tk.Label(window, text="Lien :", bg='black', fg='white', font=font_style) 
labelLink.pack(anchor="center", pady=5) 
link = tk.Entry(window, fg='white', bg='gray', font=font_style, insertbackground='white') 
link.pack(anchor="center", pady=10, ipadx=10, ipady=5) 
link.insert(0, "e.g:https://www.google.com")  
menu_bar = tk.Menu(window) 
file_menu = tk.Menu(menu_bar, tearoff=0) 
file_menu.add_command(label="Convertir le Fichier", command=openConvertFile) 
file_menu.add_command(label="Convertir le Lien", command=openConvertLink) 
file_menu.add_command(label="Recherche", command=openSearchWindow)
file_menu.add_separator() 
file_menu.add_command(label="Quitter", command=window.quit) 
help_menu = tk.Menu(menu_bar, tearoff=0) 
help_menu.add_command(label="À propos", command=showAbout) 
menu_bar.add_cascade(label="Fichier", menu=file_menu) 
menu_bar.add_cascade(label="Aide", menu=help_menu) 
window.config(menu=menu_bar)
progressLabel=tk.Label(window,text="Chargement :",fg='white',bg='black')
progressLabel.pack(anchor="center", pady=5) 
progress = ttk.Progressbar(window, orient=tk.HORIZONTAL, length=400, mode='determinate') 
progress.pack(pady=20)
window.mainloop()